import os
import unittest
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from time import sleep
from selenium import webdriver
from selenium.common.exceptions import InvalidSelectorException
import psycopg2
import json
from selenium.webdriver.support.wait import WebDriverWait
import glob


class Base:

    def get_file_dir(self, folder, filename):
        return os.path.dirname(os.path.abspath(__file__)).replace('Pages', "") + folder + "/" + filename

    def delete_all_screens(self):
        files = glob.glob('../screens/*.png')
        for f in files:
            os.remove(f)


class BasePage(unittest.TestCase):
    def __init__(self, driver: webdriver.Chrome):
        super().__init__()
        self.browser = driver

    def find_element_by_locator_type(self, locator, locator_type='css'):
        if locator_type == 'css':
            return self.browser.find_element_by_css_selector(locator)
        else:
            return self.browser.find_element_by_xpath(locator)

    def find_element_by_locator(self, locator):
        try:
            return self.browser.find_element_by_css_selector(locator)
        except InvalidSelectorException:
            return self.browser.find_element_by_xpath(locator)

    def find_elements_by_locator(self, locator):
        try:
            return self.browser.find_elements_by_css_selector(locator)
        except InvalidSelectorException:
            return self.browser.find_elements_by_xpath(locator)

    def scroll_into_view(self, locator):
        self.browser.execute_script("arguments[0].scrollIntoView();",
                                    self.find_element_by_locator(locator))

    def scroll_into_view_by_elem(self, elem):
        self.browser.execute_script("arguments[0].scrollIntoView();",
                                    elem)

    def fill_field_and_check_value_click_on_focus(self, locator, text, focus_locator, locator_type='css'):
        self.click_on_elem(locator)
        self.find_element_by_locator(locator).send_keys(text)
        self.wait_until_located(focus_locator, 'css', 6)  # todo тип фокус локатора
        self.click_on_elem(focus_locator)
        if locator_type == 'css':
            WebDriverWait(self.browser, 4).until(
                EC.text_to_be_present_in_element_value((By.CSS_SELECTOR, locator), text), 'text not found in css')
        else:
            WebDriverWait(self.browser, 4).until(
                EC.text_to_be_present_in_element_value((By.XPATH, locator), text), 'text not found in xpath')

    def fill_field_and_check_value(self, locator, text, locator_type='css'):
        self.click_on_elem(locator)
        self.find_element_by_locator(locator).send_keys(text)
        if locator_type == 'css':
            WebDriverWait(self.browser, 4).until(
                EC.text_to_be_present_in_element_value((By.CSS_SELECTOR, locator), text), 'text not found in css')
        else:
            WebDriverWait(self.browser, 4).until(
                EC.text_to_be_present_in_element_value((By.XPATH, locator), text), 'text not found in xpath')

    def fill_field(self, locator, text):
        self.click_on_elem(locator)
        self.find_element_by_locator(locator).send_keys(text)

    def fill_field_by_elem(self, elem, text):
        self.scroll_into_view_by_elem(elem)
        elem.click()
        elem.send_keys(text)

    def clear_field(self, locator):
        self.find_element_by_locator(locator).clear()

    def clear_field_by_elem(self, elem):
        elem.clear()

    def select_elem(self, locator, text, form_locator=None):
        self.click_on_elem(locator)
        if form_locator is not None:
            option = form_locator + '//option[text()[contains(.,"' + text + '")]][not(@disabled)]'
        else:
            option = '//option[text()[contains(.,"' + text + '")]][not(@disabled)]'
        self.click_on_elem(option)

    def select_elem_exact_text(self, locator, text, form_locator=None):
        self.click_on_elem(locator)
        if form_locator is not None:
            option = form_locator + "//option[text()='" + text + "'][not(@disabled)]"
        else:
            option = "//option[text()='" + text + "'][not(@disabled)]"
        self.click_on_elem(option)

    def select_by_elem(self, elem, text, form_locator=None):
        elem.click()
        if form_locator is not None:
            option = form_locator + '//option[text()[contains(.,"' + text + '")]]'
        else:
            option = '//option[text()[contains(.,"' + text + '")]]'
        self.click_on_elem(option)

    def get_field_value(self, locator):
        return self.find_element_by_locator(locator).get_attribute('value')

    def get_text(self, locator=None, elem=None):
        if locator is not None:
            return self.find_element_by_locator(locator).text
        if elem is not None:
            return elem.text

    def click_on_elem(self, locator):
        sleep(0.4)
        elem = self.find_element_by_locator(locator)
        self.scroll_into_view(locator)
        elem.click()

    def click_on_elem_with_number(self, locator, number):
        sleep(0.4)
        elem = self.find_elements_by_locator(locator)[number]
        self.scroll_into_view(locator)
        elem.click()

    def open_new_window(self):
        self.browser.execute_script("window.open('');")

    def wait_until_located(self, locator, locator_type='css', time=50):
        if locator_type == 'css':
            return WebDriverWait(self.browser, time).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, locator)))
        else:
            return WebDriverWait(self.browser, time).until(
                EC.presence_of_element_located((By.XPATH, locator)))

    def wait_until_clickable(self, locator, locator_type='css', time=30):
        if locator_type == 'css':
            return WebDriverWait(self.browser, time).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, locator)))
        else:
            return WebDriverWait(self.browser, time).until(
                EC.element_to_be_clickable((By.XPATH, locator)))

    def wait_until_disappear(self, locator, locator_type, time=10):
        if locator_type == 'css':
            WebDriverWait(self.browser, time).until(
                EC.invisibility_of_element_located((By.CSS_SELECTOR, locator)))
        else:
            WebDriverWait(self.browser, time).until(
                EC.invisibility_of_element_located((By.XPATH, locator)))

    def wait_text_to_be_present_in_element_value(self, locator, text):
        WebDriverWait(self.browser, 10).until(EC.text_to_be_present_in_element_value((By.XPATH, locator), text))

    def wait_until_alert_is_present(self, time=10):
        WebDriverWait(self.browser, time).until(EC.alert_is_present())

    def close_current_tab(self):
        self.browser.close()

    def check_selected(self, locator):
        self.assertEqual(self.find_element_by_locator(locator).is_selected(), True, msg='locator:' + locator)

    def delete_cookie(self):
        self.browser.delete_all_cookies()

    def make_settings_and_return_driver(self):
        from selenium.webdriver.chrome.options import Options
        options = Options()
        # options.add_argument('--headless')
        # options.add_argument('--disable-gpu')  # Last I checked this was necessary.
        # options.add_argument('--start-maximized')
        options.add_argument('--no-sandbox')  # Bypass OS security model
        # options.add_argument('--window-size=1920x1080')
        options.add_argument("window-size=1920,1080")
        options.add_argument('--disable-dev-shm-usage')  # overcome limited resource problems
        driver = webdriver.Chrome(options=options)
        driver.implicitly_wait(1)
        driver.set_window_size(1920, 1080)
        return driver
        # Base().delete_all_screens()

    def fill_fields(self, locator, text=None, list_texts=None, focus_locator=None):
        list_elements = self.find_elements_by_locator(locator)
        if text is not None:
            for elem in list_elements:
                self.scroll_into_view_by_elem(elem)
                elem.click()
                elem.send_keys(text)
                if focus_locator is not None:
                    self.click_on_elem(focus_locator)
        if list_texts is not None:
            for elem, text in zip(list_elements, list_texts):
                self.scroll_into_view_by_elem(elem)
                elem.click()
                elem.send_keys(text)
                if focus_locator is not None:
                    self.click_on_elem(focus_locator)

    def click_on_elements(self, locator, click_on_first=True):
        elements = self.find_elements_by_locator(locator)
        if click_on_first:
            for elem in elements:
                sleep(0.3)
                elem.click()
        else:
            for elem in elements:
                if elements.index(elem) != 0:
                    sleep(0.3)
                    elem.click()

    def compare(self, docx_file1, docx_file2):
        dir1 = Base().get_file_dir('docx', docx_file1)
        dir2 = Base().get_file_dir('docx', docx_file2)

        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import docx

        doc1 = docx.Document(dir1)
        doc2 = docx.Document(dir2)

        for paragraph_doc1, paragraph_doc2 in zip(doc1.paragraphs, doc2.paragraphs):
            self.assertEqual(paragraph_doc1.text, paragraph_doc2.text)


class Alert(BasePage):
    yes_btn = '//button[text()[contains(.,"Да")]]'

    def approve(self):
        self.click_on_elem(self.yes_btn)

    def alert_accept(self):
        self.wait_until_alert_is_present()
        alert = self.browser.switch_to.alert
        alert.accept()


class Loaders(BasePage):
    loaderAgent = "//div[contains(@class, 'loader-block')]"
    loaderBank = '.loader__block'

    def wait_until_bank_loader_disappear(self, time=10):
        self.wait_until_disappear(self.loaderBank, 'css', time)

    def wait_until_agent_loader_disappear(self, time=10):
        self.wait_until_disappear(self.loaderAgent, 'xpath', time)


class DatabaseConnect(unittest.TestCase):
    connection = None

    def __init__(self):
        super().__init__()
        with open("config_connections.json") as json_data_file:
            self.db_connection_data = json.load(json_data_file)['db']

        self.connection = psycopg2.connect(user=self.db_connection_data['user'],
                                           password=self.db_connection_data['password'],
                                           host=self.db_connection_data['host'],
                                           port=self.db_connection_data['port'],
                                           database=self.db_connection_data['database'])

        self.cursor = self.connection.cursor()

    def cleanup_data_by_inn(self, inn):
        clear_entities_query = 'delete from entities  where entity_inn=\'' + inn + '\''

        clear_individuals_query = 'delete from individuals where entity_id in' \
                                  '(select id from entities where ' \
                                  'entity_inn=\'' + inn + '\'); '

        clear_my_requests_query = 'delete from my_requests where individual_id in' \
                                  '(select id from individuals where entity_id in' \
                                  '(select id from entities where ' \
                                  'entity_inn=\'' + inn + '\')) '

        clear_customers_query = 'delete from customers where my_request_id in ' \
                                '(select id from my_requests where individual_id in' \
                                '(select id from individuals where entity_id in' \
                                '(select id from entities where entity_inn=\'' \
                                + inn + '\')));'

        clear_request_datas_query = 'delete from request_datas where my_request_id in ' \
                                    '(select id from my_requests where individual_id in' \
                                    '(select id from individuals where entity_id in' \
                                    '(select id from entities where entity_inn=\'' + inn + '\')))'

        clear_certificate_infos_query = "delete from certificate_infos where my_request_id in" \
                                        "(select my_request_id from my_requests where individual_id in" \
                                        "(select id from individuals where entity_id in" \
                                        "(select id from entities where entity_inn='" + inn + "')))"

        clear_bg_contracts_query = "delete from bg_contracts where my_request_id in(select id from my_requests " \
                                   "where individual_id in(select id from individuals where entity_id in" \
                                   "(select id from entities where entity_inn='" + inn + "')))"

        clear_bg_messages_query = "delete from bg_messages where my_request_id in " \
                                  "(select my_request_id from my_requests where individual_id in" \
                                  "(select id from individuals where entity_id in" \
                                  "(select id from entities where entity_inn='" + inn + "')))"

        clear_bg_financial_statements = "delete from financial_statements where my_request_id in " \
                                        "(select my_request_id from my_requests where individual_id in" \
                                        "(select id from individuals where entity_id in" \
                                        "(select id from entities where entity_inn='" + inn + "')))"

        clear_bg_request_approvals = "delete from request_approvals where my_request_id in " \
                                     "(select my_request_id from my_requests where individual_id in" \
                                     "(select id from individuals where entity_id in" \
                                     "(select id from entities where entity_inn='" + inn + "')))"

        clear_bg_payment_orders = "delete from payment_orders where my_request_id in " \
                                  "(select my_request_id from my_requests where individual_id in" \
                                  "(select id from individuals where entity_id in" \
                                  "(select id from entities where entity_inn='" + inn + "')))"

        clear_entity_meta_data = "delete from entity_meta_data where entity_id in " \
                                 "(select id from entities where entity_inn='" + inn + "')"

        clear_eis_query = "delete from eis where my_request_id in" \
                          "(select my_request_id from my_requests where individual_id in" \
                          "(select id from individuals where entity_id in" \
                          "(select id from entities where entity_inn='" + inn + "')))"

        my_request_id = "(select my_request_id from my_requests where individual_id in" \
                        "(select id from individuals where entity_id in" \
                        "(select id from entities where entity_inn='" + inn + "')))"

        clear_reclassifications = 'delete from reclassifications where my_request_id in' + my_request_id

        self.cursor.execute(clear_bg_contracts_query)
        self.cursor.execute(clear_certificate_infos_query)
        self.cursor.execute(clear_request_datas_query)
        # self.cursor.execute(clear_reclassifications)
        self.cursor.execute(clear_customers_query)
        self.cursor.execute(clear_bg_financial_statements)
        self.cursor.execute(clear_bg_request_approvals)
        self.cursor.execute(clear_bg_payment_orders)
        self.cursor.execute(clear_bg_messages_query)
        self.cursor.execute(clear_eis_query)
        self.cursor.execute(clear_my_requests_query)
        self.cursor.execute(clear_individuals_query)
        self.cursor.execute(clear_entity_meta_data)
        # self.cursor.execute(clear_entities_query)
        self.cursor.execute('commit')

    def change_bid_status(self, bid_id, status_text):
        status_id_query = "select id from bid_statuses b where b.title='" + status_text + "'"
        self.cursor.execute(status_id_query)
        status_id = self.cursor.fetchone()[0]
        change_status_query = "update my_requests set status=" + str(status_id) + " where id=" + str(bid_id)
        self.cursor.execute(change_status_query)
        self.cursor.execute('commit')

    def check_bid_status(self, bid_id, status_text_expected):
        get_status_text_query = 'select title from bid_statuses b ,my_requests m where m.status=b.id and m.id=' + bid_id
        self.cursor.execute(get_status_text_query)
        status_text_fact = self.cursor.fetchone()[0]
        self.assertEqual(status_text_expected, status_text_fact)
